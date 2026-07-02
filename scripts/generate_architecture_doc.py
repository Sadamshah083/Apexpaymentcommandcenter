#!/usr/bin/env python3
"""Generate ApexOne architecture documentation as DOCX and PDF."""

from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCX_PATH = DOCS / "ApexOne-Architecture.docx"
PDF_PATH = DOCS / "ApexOne-Architecture.pdf"

GENERATED = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def build_docx() -> None:
    doc = Document()
    title = doc.add_heading("ApexOne Command Center", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Mega Architecture & System Flow")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {GENERATED}")
    doc.add_paragraph("Tenant telephony: apexone.morpheus.cx")
    doc.add_page_break()

    # --- Executive summary ---
    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "ApexOne is a multi-tenant Laravel application that serves as the command center for sales "
        "operations, lead pipeline management, email tooling, and communications. Morpheus CX "
        "(apexone.morpheus.cx) is the telephony backend. The app exposes two web portals — Admin "
        "(/admin) and Agent (/portal) — scoped per Workspace with role-based access.",
    )

    add_heading(doc, "1.1 One-line summary", 2)
    add_para(
        doc,
        "Admins import and enrich leads, assign them to setters/closers through role-based portals, "
        "and agents dial via Morpheus (REST for call control, SIP for outbound) — all scoped per "
        "workspace with optional module permissions and local call-history logging where Morpheus "
        "has no history API.",
    )

    # --- Tech stack ---
    add_heading(doc, "2. Technology Stack", 1)
    add_table(
        doc,
        ["Layer", "Technology"],
        [
            ["Application", "Laravel 13 (PHP), Blade templates, Hotwire Turbo"],
            ["Frontend assets", "Vite, Tailwind CSS 4"],
            ["Morpheus integration", "REST API (JSON + X-API-Key), Laravel Http client"],
            ["Outbound voice", "SIP softphone (Zoiper/Linphone) — not REST originate"],
            ["Realtime UI", "HTTP polling (workspace_sync_events)"],
            ["AI enrichment", "Google Gemini, OpenRouter (queued jobs)"],
            ["Email tooling", "DNS/MX/SMTP verification, IMAP inbox tests"],
            ["Data", "SQL database, workspace-scoped Eloquent models"],
            ["Push", "Web Push (VAPID) via service worker"],
        ],
    )

    # --- Portals ---
    add_heading(doc, "3. Portals & Authentication", 1)
    add_table(
        doc,
        ["Portal", "URL", "Who", "Middleware"],
        [
            ["Admin Portal", "/admin/*", "Super Admin, Admin, Manager", "AdminPortalMiddleware + EnsureAdminModuleAccess"],
            ["Agent Portal", "/portal/*", "Setter, Setter TL, Closer, Closer TL", "MarketerPortalMiddleware"],
            ["Public", "/", "Redirects to portal login", "—"],
        ],
    )
    add_heading(doc, "3.1 Auth flows", 2)
    add_bullets(
        doc,
        [
            "Register workspace: /admin/register-workspace → creates User + Workspace (owner = super_admin).",
            "Admin login: /admin/login (username + password) → module-gated admin dashboard.",
            "Portal login: /portal/login → role-based dashboard (setter/closer/team lead).",
            "Member creation: Super Admin creates agents with username/password (no email invite flow).",
            "Single User table — portal vs admin is determined by role on workspace_user pivot.",
        ],
    )

    # --- Roles ---
    add_heading(doc, "4. User Roles & Access", 1)
    add_table(
        doc,
        ["Role", "Portal", "Primary access"],
        [
            ["super_admin", "Admin", "Full access; workspace owner (workspaces.admin_id)"],
            ["admin", "Admin", "Modules granted via module_permissions"],
            ["manager", "Admin", "Modules granted via module_permissions"],
            ["appointment_setter_team_lead", "Portal", "Team setter queue, pipeline oversight"],
            ["appointment_setter", "Portal", "Own assigned leads, dial via Communications Hub"],
            ["closers_team_lead", "Portal", "Assign closers, appointment_settled queue"],
            ["closer", "Portal", "Own closer-assigned leads, close deals"],
        ],
    )
    add_heading(doc, "4.1 Admin modules (config/admin_modules.php)", 2)
    add_table(
        doc,
        ["Module", "Purpose"],
        [
            ["dashboard", "Overview metrics"],
            ["lead_pipeline", "Workflows, lead approve/reject"],
            ["lead_tags", "Tag batch enrich/distribute"],
            ["email_lists", "Bulk email verifier"],
            ["deliverability", "Domain deliverability scan"],
            ["content_analyzer", "Outbound spam analyzer"],
            ["reputation", "Sender reputation center"],
            ["sales_ops", "Performance, distribution, reactivation"],
            ["crm", "CRM campaigns with AI research"],
            ["business_research", "Standalone AI company research"],
            ["communications", "Communications Hub (Morpheus)"],
            ["server_monitoring", "Queue/CPU monitoring (super_admin only)"],
            ["user_management", "Workspace members (super_admin only)"],
        ],
    )

    # --- Workspace ---
    add_heading(doc, "5. Workspace Multi-Tenancy", 1)
    add_para(doc, "Every customer organization is a Workspace. All data is scoped by workspace_id.")
    add_table(
        doc,
        ["Concept", "Details"],
        [
            ["Workspace", "Tenant boundary: workflows, leads, CRM, call logs, email lists"],
            ["workspace_user pivot", "role, status, module_permissions, morpheus_user_id, morpheus_extension_num"],
            ["current_workspace_id", "Active tenant context on User model"],
            ["WorkspaceContextService", "Resolves active workspace for authenticated user"],
            ["WorkspaceSyncService", "Records events; clients poll for live UI updates"],
        ],
    )

    # --- Morpheus ---
    add_heading(doc, "6. Morpheus CX Integration", 1)
    add_heading(doc, "6.1 Why Morpheus", 2)
    add_bullets(
        doc,
        [
            "Replaced Zoom Phone as the telephony backend.",
            "Provides live call control, queues, conferences, dialer leads/campaigns.",
            "ApexOne is the UI + workspace layer; Morpheus is the PBX/call engine.",
        ],
    )
    add_heading(doc, "6.2 API endpoints used", 2)
    add_table(
        doc,
        ["API", "Base URL", "Auth"],
        [
            ["Call-Control (primary)", "https://apexone.morpheus.cx/api/v1/call-control", "X-API-Key (MORPHEUS_API_KEY)"],
            ["Platform (optional)", "https://apexone.morpheus.cx/api/v1/*", "MORPHEUS_PLATFORM_API_KEY"],
        ],
    )
    add_heading(doc, "6.3 Architecture layers (Communications)", 2)
    add_code(
        doc,
        ".env (MORPHEUS_HOST, MORPHEUS_API_KEY, SIP settings)\n"
        "    ↓\n"
        "ZoomApiService — HTTP REST client (legacy name from Zoom migration)\n"
        "    ↓\n"
        "MorpheusHubService — cached reads for hub panels\n"
        "CommunicationsDataService / CommunicationsInboxService — inbox UI data\n"
        "    ↓\n"
        "MorpheusHubController — POST/PATCH/DELETE (call control, CRUD)\n"
        "CommunicationsHubController — Communications Hub pages\n"
        "    ↓\n"
        "Blade UI (calls, dialer, queues, team, phone agents, etc.)",
    )
    add_heading(doc, "6.4 Call-Control API coverage", 2)
    add_table(
        doc,
        ["Section", "Client", "Hub UI"],
        [
            ["Calls (list/get)", "Yes", "Calls channel — live calls only"],
            ["Call actions", "Yes", "Hold, transfer, park, bridge, disposition"],
            ["Queues", "Full CRUD", "Queues panel"],
            ["Conferences", "Full CRUD + members", "Conferences panel"],
            ["Campaigns", "Full CRUD", "Campaigns panel"],
            ["Lists", "Full CRUD", "Lists panel"],
            ["Leads", "Full CRUD", "Leads panel"],
            ["Users", "Full CRUD", "Team panel"],
            ["Extensions", "Full CRUD", "Extensions + Phone agents"],
            ["Originate", "Probed — API returns 405", "SIP softphone dialer instead"],
            ["Historical CDR", "Not in API", "Local communication_call_logs table"],
        ],
    )
    add_heading(doc, "6.5 Agent provisioning flow", 2)
    add_bullets(
        doc,
        [
            "Admin creates workspace user in User Management.",
            "Admin provisions in Communications → Phone agents (extension + SIP password).",
            "Agent registers softphone to apexone.morpheus.cx.",
            "Agent dials from hub → SIP launch page opens call.",
        ],
    )

    # --- Lead pipeline ---
    add_heading(doc, "7. Lead Pipeline Lifecycle", 1)
    add_para(doc, "Pipeline phases (config/sales_ops.pipeline_phases):")
    add_code(
        doc,
        "imported → enriching → enriched → with_setter → appointment_settled → with_closer → closed",
    )
    add_heading(doc, "7.1 Admin flow", 2)
    add_bullets(
        doc,
        [
            "Upload CSV/XLS → WorkflowService.createFromUpload",
            "AI column mapping (WorkflowAiMapper / Gemini)",
            "ProcessWorkflowJob ingests rows → WorkflowLead records",
            "Optional: ProcessLeadJob enriches each lead (WorkflowExtractor)",
            "SetterDistributionService round-robin OR manual admin distribute",
            "Lead enters pipeline_phase: with_setter",
        ],
    )
    add_heading(doc, "7.2 Setter flow (Portal)", 2)
    add_bullets(
        doc,
        [
            "Setter Dashboard shows assigned leads",
            "Dial via Communications Hub → SIP → Morpheus live call",
            "Update setter_status; log LeadActivity",
            "On appointment_settled → handoff to Closer Team Lead",
        ],
    )
    add_heading(doc, "7.3 Closer flow (Portal)", 2)
    add_bullets(
        doc,
        [
            "Closer Team Lead assigns closer from appointment_settled queue",
            "Closer works lead in with_closer phase",
            "Closer records sale_made or closed_lost → phase: closed",
        ],
    )

    # --- Modules ---
    add_heading(doc, "8. Major Modules", 1)
    add_table(
        doc,
        ["Module", "Controller / Service", "Purpose"],
        [
            ["Workflows", "WorkflowController, WorkflowService", "CSV import, AI enrich, distribute to setters"],
            ["Pipeline", "PipelineController, LeadPipelineService", "Role dashboards for setters/closers"],
            ["Sales Ops", "SalesOpsController, SdrPerformanceService", "Performance, reactivation, activity"],
            ["CRM", "CrmCampaignController", "Parallel CSV + per-lead AI research"],
            ["Email Lists", "EmailListController", "Bulk email verification"],
            ["Deliverability", "DeliverabilityController", "Domain auth + inbox tests"],
            ["Content Analyzer", "ContentAnalyzerController", "Spam/score analysis"],
            ["Reputation", "ReputationController", "Warmup and compliance"],
            ["Business Research", "BusinessResearchController", "Standalone AI research jobs"],
            ["Communications", "CommunicationsHubController, MorpheusHubController", "Unified comms hub"],
            ["Lead Tags", "LeadTagController", "Tag-based batch operations"],
            ["Server Monitoring", "ServerMonitoringController", "Queue/CPU (super_admin)"],
        ],
    )

    # --- External ---
    add_heading(doc, "9. External Systems", 1)
    add_table(
        doc,
        ["System", "Used for"],
        [
            ["Morpheus CX", "Telephony: calls, queues, extensions, dialer data"],
            ["Google Gemini", "Lead enrichment, business research, CRM research"],
            ["OpenRouter", "Fallback AI provider"],
            ["DuckDuckGo", "Web search context for enrichment"],
            ["DNS / MX / SMTP", "Email verification and deliverability"],
            ["IMAP", "Deliverability inbox placement tests"],
            ["Web Push (VAPID)", "Browser notifications"],
        ],
    )

    # --- Data flows diagram text ---
    add_heading(doc, "10. System Flow Diagram (Mermaid)", 1)
    add_para(doc, "Paste into mermaid.live or any Mermaid renderer for visual diagram.")
    add_code(
        doc,
        MERMAID_BIG,
    )

    add_heading(doc, "11. Lead Flow Diagram (Mermaid)", 1)
    add_code(doc, MERMAID_LEAD)

    add_heading(doc, "12. Morpheus Sequence Diagram (Mermaid)", 1)
    add_code(doc, MERMAID_SEQ)

    add_heading(doc, "13. Key File Paths", 1)
    add_table(
        doc,
        ["Area", "Path"],
        [
            ["Routes", "routes/web.php, routes/morpheus-communications.php"],
            ["Auth", "app/Http/Controllers/WorkspaceAuthController.php"],
            ["User / Workspace", "app/Models/User.php, app/Models/Workspace.php"],
            ["Role config", "config/sales_ops.php"],
            ["Module config", "config/admin_modules.php"],
            ["Morpheus config", "config/integrations.php"],
            ["Morpheus API client", "app/Services/Integrations/ZoomApiService.php"],
            ["Platform API", "app/Services/Integrations/MorpheusPlatformApiService.php"],
            ["Comms hub", "app/Http/Controllers/CommunicationsHubController.php"],
            ["Call actions", "app/Http/Controllers/MorpheusHubController.php"],
            ["Workflow orchestration", "app/Services/Workflow/WorkflowService.php"],
            ["Setter distribution", "app/Services/Pipeline/SetterDistributionService.php"],
            ["Admin sidebar", "resources/views/layouts/partials/sidebar-nav-admin.blade.php"],
            ["Portal sidebar", "resources/views/layouts/partials/sidebar-nav-portal.blade.php"],
        ],
    )

    add_heading(doc, "14. Environment Variables (Morpheus)", 1)
    add_code(
        doc,
        "MORPHEUS_HOST=apexone.morpheus.cx\n"
        "MORPHEUS_API_KEY=ck_...\n"
        "MORPHEUS_DIAL_METHOD=sip\n"
        "MORPHEUS_SIP_PARAMS=user=phone\n"
        "MORPHEUS_PORTAL_URL=https://apexone.morpheus.cx/\n"
        "MORPHEUS_PLATFORM_API_KEY=  # optional, for recordings/chat",
    )

    DOCS.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    print(f"DOCX: {DOCX_PATH}")


MERMAID_BIG = """flowchart TB
    subgraph PEOPLE["People and Roles"]
        SA["Super Admin"]
        AD["Admin / Manager"]
        STL["Setter Team Lead"]
        ST["Appointment Setter"]
        CTL["Closer Team Lead"]
        CL["Closer"]
    end
    subgraph BROWSER["Browser Clients"]
        ADMIN_UI["Admin Portal /admin"]
        PORTAL_UI["Agent Portal /portal"]
    end
    subgraph CORE["Laravel Monolith"]
        ADMIN_MODS["Admin Modules"]
        PORTAL_MODS["Portal Modules"]
        SERVICES["Domain Services"]
        JOBS["Queue Workers"]
        DATA["Workspace-Scoped DB"]
    end
    subgraph EXTERNAL["External Systems"]
        MOR["Morpheus CX REST API"]
        SIP["SIP Softphone"]
        GEM["Google Gemini"]
        DNS["DNS/SMTP/IMAP"]
    end
    SA --> ADMIN_UI
    AD --> ADMIN_UI
    STL --> PORTAL_UI
    ST --> PORTAL_UI
    CTL --> PORTAL_UI
    CL --> PORTAL_UI
    ADMIN_UI --> ADMIN_MODS --> SERVICES --> DATA
    PORTAL_UI --> PORTAL_MODS --> SERVICES
    SERVICES --> JOBS
    SERVICES --> MOR
    ST --> SIP --> MOR
    SERVICES --> GEM
    SERVICES --> DNS"""


MERMAID_LEAD = """flowchart TD
    UP["Upload CSV"] --> MAP["AI Column Mapping"]
    MAP --> IMP["Import ProcessWorkflowJob"]
    IMP --> ENR["Enrich ProcessLeadJob"]
    ENR --> RR["Round-robin to Setter"]
    RR --> DIAL["Setter dials via Hub"]
    DIAL --> STATUS["Update setter_status"]
    STATUS --> HANDOFF["Handoff to Closer TL"]
    HANDOFF --> CLOSE["Closer closes deal"]"""


MERMAID_SEQ = """sequenceDiagram
    participant Admin
    participant Hub as ApexOne Hub
    participant Agent
    participant Phone as SIP Softphone
    participant API as Morpheus API
    Admin->>Hub: Provision Phone Agent
    Hub->>API: POST /users, /extensions
    Agent->>Hub: Open dialer
    Hub->>Agent: SIP launch page
    Agent->>Phone: Place call
    Hub->>API: GET /calls, POST disposition"""


def _pdf_safe(text: str) -> str:
    return (
        text.replace("\u2014", "-")
        .replace("\u2013", "-")
        .replace("\u2192", "->")
        .replace("\u2605", "*")
        .replace("\u2022", "-")
        .encode("latin-1", errors="replace")
        .decode("latin-1")
    )


def build_pdf() -> None:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 12, "ApexOne Command Center", ln=True, align="C")
    pdf.set_font("Helvetica", "", 12)
    pdf.cell(0, 8, "Mega Architecture & System Flow", ln=True, align="C")
    pdf.cell(0, 8, f"Generated: {GENERATED}", ln=True, align="C")
    pdf.ln(6)

    sections: list[tuple[str, str]] = [
        (
            "1. Executive Summary",
            "ApexOne is a multi-tenant Laravel application: Admin Portal (/admin) and Agent Portal "
            "(/portal), scoped per Workspace. Morpheus CX (apexone.morpheus.cx) is the telephony "
            "backend. Admins manage leads, email tools, and communications; agents work the "
            "pipeline and dial via SIP + Morpheus REST API.",
        ),
        (
            "2. Technology Stack",
            "Laravel 13, Blade, Turbo, Vite, Tailwind. Morpheus: REST JSON + API key. Outbound: "
            "SIP softphone. AI: Gemini/OpenRouter via queue jobs. Realtime: HTTP polling. "
            "No FastAPI, GraphQL, or WebSockets.",
        ),
        (
            "3. Portals",
            "Admin (/admin): Super Admin, Admin, Manager - module permissions via workspace_user. "
            "Portal (/portal): Setter, Setter TL, Closer, Closer TL - pipeline dashboards. "
            "Single User table; role on pivot determines portal access.",
        ),
        (
            "4. User Roles",
            "super_admin: full access, workspace owner. admin/manager: granted modules. "
            "appointment_setter: own leads + dial. appointment_setter_team_lead: team queue. "
            "closers_team_lead: assign closers. closer: close deals.",
        ),
        (
            "5. Workspace Tenancy",
            "Workspace = tenant. workspace_user pivot: role, status, module_permissions, "
            "morpheus_extension_num. All WorkflowLeads, CRM, call logs scoped by workspace_id.",
        ),
        (
            "6. Morpheus Integration",
            "Call-Control API: /api/v1/call-control with X-API-Key. ZoomApiService is the HTTP "
            "client. Outbound: SIP (originate API returns 405). Local call history fills CDR gap. "
            "Agent flow: Admin provisions extension -> softphone register -> hub dialer -> SIP.",
        ),
        (
            "7. Lead Pipeline",
            "Phases: imported -> enriching -> enriched -> with_setter -> appointment_settled -> "
            "with_closer -> closed. Admin uploads CSV, AI maps columns, jobs enrich, round-robin "
            "to setters. Setters dial and update status. Closer TL assigns closer. Closer closes.",
        ),
        (
            "8. Major Modules",
            "Workflows, Pipeline, Sales Ops, CRM, Email Lists, Deliverability, Content Analyzer, "
            "Reputation, Business Research, Communications Hub, Lead Tags, Server Monitoring.",
        ),
        (
            "9. External Systems",
            "Morpheus CX, Google Gemini, OpenRouter, DuckDuckGo, DNS/MX/SMTP, IMAP, Web Push.",
        ),
        (
            "10. Key Files",
            "routes/web.php, routes/morpheus-communications.php, ZoomApiService.php, "
            "MorpheusHubController.php, CommunicationsHubController.php, WorkflowService.php, "
            "config/integrations.php, config/sales_ops.php, config/admin_modules.php.",
        ),
    ]

    for title, body in sections:
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 9, _pdf_safe(title), ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5, _pdf_safe(body))
        pdf.ln(3)

    pdf.add_page()
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 9, "Mermaid Diagrams (paste into mermaid.live)", ln=True)
    pdf.set_font("Courier", "", 7)
    for label, diagram in [
        ("System overview", MERMAID_BIG),
        ("Lead flow", MERMAID_LEAD),
        ("Morpheus sequence", MERMAID_SEQ),
    ]:
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 7, _pdf_safe(label), ln=True)
        pdf.set_font("Courier", "", 7)
        pdf.multi_cell(0, 3.5, _pdf_safe(diagram))
        pdf.ln(2)

    DOCS.mkdir(parents=True, exist_ok=True)
    pdf.output(str(PDF_PATH))
    print(f"PDF:  {PDF_PATH}")


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print("Done.")
